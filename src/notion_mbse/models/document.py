#!/usr/bin/env python
"""
@package   document
Details:
Created:   Sunday, June 30th 2024, 8:36:38 pm
-----
Last Modified: 07/07/2024 07:07:57
Modified By: Mathew Cosgrove
-----
"""

__author__ = "Mathew Cosgrove"
__file__ = "document.py"
__version__ = "0.1.0"
import io
from enum import Enum
from pathlib import Path
from typing import List
from typing import Optional

import markdown
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

# from pydantic.schema import schema
from pydantic import Field

from .element import Element


class DocumentType(Enum):
    """An enumeration of the different types of documents."""

    GENERAL = "General"
    REQUIREMENTS = "Requirements"
    DESIGN = "Design"
    IMPLEMENTATION = "Implementation"
    TEST = "Test"
    DEPLOYMENT = "Deployment"
    OPERATION = "Operation"
    MAINTENANCE = "Maintenance"
    SUPPORT = "Support"
    TRAINING = "Training"
    USER_MANUAL = "User Manual"
    TECHNICAL_MANUAL = "Technical Manual"
    PRESENTATION = "Presentation"
    REPORT = "Report"
    PROPOSAL = "Proposal"
    SPECIFICATION = "Specification"
    PLAN = "Plan"


class DocumentSectionType(Enum):
    """An enumeration of the different types of document sections."""

    OVERVIEW = "Overview"
    INTRODUCTION = "Introduction"
    BACKGROUND = "Background"
    SCOPE = "Scope"
    OBJECTIVES = "Objectives"
    REQUIREMENTS = "Requirements"
    DESIGN_OVERVIEW = "Design Overview"
    ARCHITECTURE = "Architecture"
    COMPONENTS = "Components"
    INTERFACES = "Interfaces"
    DATA = "Data"
    ALGORITHMS = "Algorithms"
    TEST_PLAN = "Test Plan"
    TEST_CASES = "Test Cases"
    TEST_RESULTS = "Test Results"
    DEPLOYMENT_PLAN = "Deployment Plan"
    DEPLOYMENT_RESULTS = "Deployment Results"
    OPERATION_PLAN = "Operation Plan"
    OPERATION_RESULTS = "Operation Results"
    MAINTENANCE_PLAN = "Maintenance Plan"
    MAINTENANCE_RESULTS = "Maintenance Results"
    SUPPORT_PLAN = "Support Plan"


class DocumentSectionFormatType(Enum):
    SLIDE = "Slide"
    PAGE = "Page"
    CHAPTER = "Chapter"
    SECTION = "Section"
    SUBSECTION = "Subsection"
    PARAGRAPH = "Paragraph"
    BULLET = "Bullet"
    NUMBERED = "Numbered"
    TABLE = "Table"
    FIGURE = "Figure"
    EQUATION = "Equation"
    CODE = "Code"
    LIST = "List"
    REFERENCE = "Reference"
    FOOTNOTE = "Footnote"
    ENDNOTE = "Endnote"
    GLOSSARY = "Glossary"
    INDEX = "Index"
    APPENDIX = "Appendix"
    ANNEX = "Annex"
    ABBREVIATION = "Abbreviation"
    ACRONYM = "Acronym"
    SYMBOL = "Symbol"
    KEYWORD = "Keyword"
    SUMMARY = "Summary"


class DocumentSection(Element):
    type: Optional[DocumentSectionType] = Field(None, description="The type of the section.")
    sub_type: Optional[DocumentSectionFormatType] = Field(None, description="The format of the section.")
    section_level: Optional[int] = Field(None, description="The level of the section.")
    section_number: Optional[int] = Field(None, description="The number of the section.")
    document: Optional[str] = Field(None, description="The document the section belongs to.")
    page: Optional[int] = Field(None, description="The page the section belongs to.")
    parent_id: Optional[str] = Field(None, description="The parent section id.")
    author: Optional[str] = Field(None, description="The author of the section.")
    content: Optional[str] = Field(None, description="The content of the section.")
    section_length: Optional[int] = Field(None, description="The length of the section.")
    references: Optional[List[str]] = Field([], description="The references of the section.")
    summary: Optional[str] = Field(None, description="The summary of the section.")
    keywords: Optional[List[str]] = Field([], description="The keywords of the section.")
    related_sections: Optional[List[str]] = Field([], description="The related sections of the section.")


def html_to_docsections(html_content: str, document_name: str) -> List[DocumentSection]:
    soup = BeautifulSoup(html_content, "html.parser")
    sections = []
    section_number = 0

    for header in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        section_number += 1
        section_level = int(header.name[1])
        section_name = header.get_text()

        # Assuming section description might be in the next sibling paragraph
        section_description = None
        next_sibling = header.find_next_sibling()
        if next_sibling and next_sibling.name == "p":
            section_description = next_sibling.get_text()

        section_content = ""
        next_sibling = header.find_next_sibling()
        while next_sibling and next_sibling.name not in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            section_content += str(next_sibling)
            next_sibling = next_sibling.find_next_sibling()

        section = DocumentSection(
            name=section_name,
            section_level=section_level,
            section_description=section_description,
            document=document_name,
            section_number=section_number,
            content=section_content,
        )
        sections.append(section)

    return sections


class Document(Element):
    type: Optional[DocumentType] = Field(None, description="The type of the document.")
    section_ids: Optional[List[str]] = Field([], description="The section ids of the document.")
    summary: Optional[str] = Field(None, description="The summary of the document.")
    keywords: Optional[List[str]] = Field([], description="The keywords of the document.")
    references: Optional[List[str]] = Field([], description="The references of the document.")
    sections: Optional[List[DocumentSection]] = Field([], description="The sections of the document.")

    def add_section(self, section: DocumentSection):
        self.sections.append(section)
        self.section_ids.append(section.id)

    def remove_section(self, section_id: str):
        self.sections = [section for section in self.sections if section.id != section_id]
        self.section_ids = [section_id for section_id in self.section_ids if section_id != section_id]

    def load_from_html(self, html_content: str):
        self.sections = html_to_docsections(html_content, self.name)
        for section in self.sections:
            self.section_ids.append(section.id)

    def load_from_markdown(self, markdown_content: str):
        # Convert Markdown to HTML
        html_content = markdown.markdown(markdown_content)
        self.load_from_html(html_content)

    def load_from_docx(self, docx_file: str):
        html_content = DocxDocument(docx_file)
        self.load_from_html(html_content)

    def load_from_docx_file(self, docx_file: str):
        doc = DocxDocument(docx_file)
        current_section = None
        for paragraph in doc.paragraphs:
            print(paragraph.style.name)
            if paragraph.style.name.startswith("Heading"):
                if current_section:
                    self.sections.append(current_section)
                    self.section_ids.append(current_section.id)

                print(f"Section Change: {paragraph.text}")
                ## TODO: Determine level based on something
                level = 0
                current_section = DocumentSection(
                    name=paragraph.text,
                    section_level=level,
                    content="",
                    document=self.name,
                )
            elif current_section:
                current_section.content += paragraph.text + "\n"

        if current_section:
            self.sections.append(current_section)
            self.section_ids.append(current_section.id)

    def load_from_docx_direct_simple(self, docx_content: bytes):
        doc = DocxDocument(io.BytesIO(docx_content))
        for paragraph in doc.paragraphs:
            section = DocumentSection(
                name=paragraph.text,
                content=paragraph.text,
                document=self.name,
            )
            self.sections.append(section)
            self.section_ids.append(section.id)

    def load_from_pdf_simple(self, pdf_content: bytes):
        from PyPDF2 import PdfReader

        pdf_reader = PdfReader(io.BytesIO(pdf_content))
        for page in pdf_reader.pages:
            text = page.extract_text()
            section = DocumentSection(
                name=f"Page {pdf_reader.pages.index(page) + 1}",
                content=text,
                document=self.name,
            )
            self.sections.append(section)
            self.section_ids.append(section.id)

    def load_from_docx_direct(self, docx_content: bytes):
        doc = DocxDocument(io.BytesIO(docx_content))
        current_section = None
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                if current_section:
                    self.sections.append(current_section)
                    self.section_ids.append(current_section.id)
                level = int(paragraph.style.name[-1])
                current_section = DocumentSection(
                    name=paragraph.text,
                    section_level=level,
                    content="",
                    document=self.name,
                )
            elif current_section:
                current_section.content += paragraph.text + "\n"

        if current_section:
            self.sections.append(current_section)
            self.section_ids.append(current_section.id)

    def load_from_pdf(self, pdf_content: bytes):
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTChar
        from pdfminer.layout import LTTextContainer

        pdf_file = io.BytesIO(pdf_content)
        current_section = None
        for page_layout in extract_pages(pdf_file):
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    text = element.get_text().strip()
                    if text:
                        font_sizes = [char.size for char in element if isinstance(char, LTChar)]
                        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 0

                        # Assume larger font sizes indicate headers/section titles
                        if avg_font_size > 12 and len(text) < 100:  # Adjust these thresholds as needed
                            if current_section:
                                self.sections.append(current_section)
                                self.section_ids.append(current_section.id)
                            current_section = DocumentSection(
                                name=text,
                                section_level=1,  # You might want to adjust this based on font size
                                content="",
                                document=self.name,
                            )
                        elif current_section:
                            current_section.content += text + "\n"

        if current_section:
            self.sections.append(current_section)
            self.section_ids.append(current_section.id)

    def load(self, file_path: str, simple: bool = False):
        if file_path.endswith(".html"):
            with Path.open(file_path) as file:
                self.load_from_html(file.read())
        elif file_path.endswith(".md"):
            with Path.open(file_path) as file:
                self.load_from_markdown(file.read())
        elif file_path.endswith(".docx"):
            if simple:
                with Path.open(file_path, "rb") as file:
                    self.load_from_docx_direct_simple(file.read())
            else:
                self.load_from_docx(file_path)
        elif file_path.endswith(".pdf"):
            if simple:
                with Path.open(file_path, "rb") as file:
                    self.load_from_pdf_simple(file.read())
            else:
                self.load_from_pdf(file_path)
        else:
            raise ValueError("Unsupported file format")

    def load_from_url(self, url: str, simple: bool = False):
        response = requests.get(url, timeout=10)
        content = response.content
        if url.endswith(".html"):
            self.load_from_html(content.decode("utf-8"))
        elif url.endswith(".md"):
            self.load_from_markdown(content.decode("utf-8"))
        elif url.endswith(".docx"):
            if simple:
                self.load_from_docx_direct_simple(content)
            else:
                self.load_from_docx_direct(content)
        elif url.endswith(".pdf"):
            if simple:
                self.load_from_pdf_simple(content)
            else:
                self.load_from_pdf(content)
        else:
            raise ValueError("Unsupported file format")

    def export_to_html(self) -> str:
        html = "<html><body>"
        for section in self.sections:
            html += f"<h2>{section.name}</h2>"
            html += f"<div>{section.content}</div>"
        html += "</body></html>"
        return html

    def export_to_docx(self) -> bytes:
        doc = DocxDocument()
        for section in self.sections:
            doc.add_heading(section.name, level=2)
            doc.add_paragraph(section.content)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_to_pdf(self) -> bytes:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        for section in self.sections:
            c.setFont("Helvetica-Bold", 14)
            c.drawString(72, height - 72, section.name)
            c.setFont("Helvetica", 12)
            text = c.beginText(72, height - 100)
            text.textLines(section.content)
            c.drawText(text)
            c.showPage()
        c.save()
        buffer.seek(0)
        return buffer.getvalue()

    def export_to_markdown(self) -> str:
        markdown_content = ""
        for section in self.sections:
            level = "#" * section.section_level
            markdown_content += f"{level} {section.name}\n\n"
            markdown_content += f"{section.content}\n\n"
        return markdown_content

    def export(self, file_path: str):
        if file_path.endswith(".html"):
            with Path.open(file_path, "w") as file:
                file.write(self.export_to_html())
        elif file_path.endswith(".docx"):
            with Path.open(file_path, "wb") as file:
                file.write(self.export_to_docx())
        elif file_path.endswith(".pdf"):
            with Path.open(file_path, "wb") as file:
                file.write(self.export_to_pdf())
        else:
            raise ValueError("Unsupported file format")


def load_document(file_path: str, simple: bool = False) -> Document:
    doc = Document(name=file_path.split("/")[-1])
    doc.load(file_path, simple)
    return doc


def load_document_from_url(url: str, simple: bool = False) -> Document:
    doc = Document(name=url)
    doc.load_from_url(url, simple)
    return doc


class FileType(Enum):
    """An enumeration of the different types of files."""

    TEXT = "Text"
    IMAGE = "Image"
    AUDIO = "Audio"
    VIDEO = "Video"
    DOCUMENT = "Document"
    SPREADSHEET = "Spreadsheet"
    PRESENTATION = "Presentation"
    DATABASE = "Database"
    CODE = "Code"
    EXECUTABLE = "Executable"
    ARCHIVE = "Archive"
    COMPRESSED = "Compressed"
    DISK_IMAGE = "Disk Image"
    VIRTUAL_MACHINE = "Virtual Machine"
    CONFIGURATION = "Configuration"
    LOG = "Log"
    BACKUP = "Backup"
    TEMPORARY = "Temporary"
    SYSTEM = "System"
    WEB = "Web"
    EMAIL = "Email"
    CONTACT = "Contact"
    CALENDAR = "Calendar"
    NOTE = "Note"
    BOOKMARK = "Bookmark"
    LINK = "Link"
    REFERENCE = "Reference"
    TEMPLATE = "Template"
    FORM = "Form"
    SURVEY = "Survey"
    QUESTIONNAIRE = "Questionnaire"
    POLL = "Poll"
    QUIZ = "Quiz"
    EXAM = "Exam"
    TEST = "Test"
    SURVEILLANCE = "Surveillance"
    MONITORING = "Monitoring"
    AUDIT = "Audit"
    INSPECTION = "Inspection"
    REVIEW = "Review"
    EVALUATION = "Evaluation"
    ASSESSMENT = "Assessment"
    ANALYSIS = "Analysis"
    REPORT = "Report"
    PROPOSAL = "Proposal"
    SPECIFICATION = "Specification"
    PLAN = "Plan"
    AGREEMENT = "Agreement"
    CONTRACT = "Contract"
    LICENSE = "License"
    CERTIFICATE = "Certificate"
    DIPLOMA = "Diploma"
    DEGREE = "Degree"


class FileDocument(Document):
    type: Optional[FileType] = Field(None, description="The type of the file.")
    file_path: Optional[str] = Field(None, description="The path of the file.")
    file_size: Optional[int] = Field(None, description="The size of the file.")
    file_format: Optional[str] = Field(None, description="The format of the file.")
    file_ext: Optional[str] = Field(None, description="The extension of the file.")

    def load(self, file_path: str, simple: bool = False):
        import mimetypes

        self.file_path = file_path
        self.file_size = Path.stat(file_path).st_size
        self.file_format = mimetypes.guess_type(file_path)[0]
        self.file_ext = Path.suffix(file_path)[1][1:]
        if self.file_ext in ["html", "docx", "pdf"]:
            self.load(file_path, simple)
        else:
            raise ValueError("Unsupported file format")


# class Specification(BaseModel):
#     model_config = ConfigDict(arbitrary_types_allowed=True)
#     name: str
#     acronym: Optional[str] = None
#     version: str
#     status: str
#     reference_id: Optional[str] = None
#     published_date: datetime.date
#     link: str


# class Specification(Document):
#     name: str
#     acronym: Optional[str] = None
#     version: str
#     status: str
#     reference_id: Optional[str] = None
#     published_date: datetime.date
#     link: str

#     @classmethod
#     def from_specification(cls, spec: Specification):
#         init_dict = {
#             "name": spec.acronym,
#             "spec_name": spec.name,
#             "version": spec.version,
#             "status": spec.status,
#             "published_date": spec.published_date,
#             "link": spec.link,
#         }
#         return cls.new(**init_dict)

#     def __str__(self):
#         return f"Specification({self.name}, {self.spec_name}, {self.link})"

#     def model_dump(self):
#         return json.dumps(self.dict(), indent=4)
